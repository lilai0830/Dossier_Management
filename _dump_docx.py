import docx

def dump(f, label):
    print("=" * 78)
    print("FILE:", f, "(" + label + ")")
    print("=" * 78)
    d = docx.Document(f)
    for i, p in enumerate(d.paragraphs):
        t = p.text.strip()
        if not t:
            continue
        print(f"[{i}] ({p.style.name}) {t}")
    for ti, tbl in enumerate(d.tables):
        print(f"--- TABLE {ti} (rows={len(tbl.rows)}, cols={len(tbl.columns)}) ---")
        for r in tbl.rows:
            print(" | ".join(c.text.strip() for c in r.cells))

dump("companion/output_framework.docx", "OUTPUT FRAMEWORK")
print("\n\n")
dump("companion/reference_library.docx", "REFERENCE LIBRARY")
